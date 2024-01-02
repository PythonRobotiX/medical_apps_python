import datetime
import calendar
import tkinter as tk
from tkinter import ttk
from tkinter import simpledialog, filedialog
from tkcalendar import DateEntry
from playsound import playsound
import time
import pandas as pd
from reportlab.pdfgen import canvas
from docx import Document

class PatientEntryList:
    def __init__(self):
        self.entries = pd.DataFrame(columns=['Patient Name', 'Entry Time'])

    def append_entry(self, patient_name):
        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        entry_data = {'Patient Name': patient_name, 'Entry Time': current_time}
        self.entries = pd.concat([self.entries, pd.DataFrame([entry_data])], ignore_index=True)

    def delete_entry(self, index):
        self.entries = self.entries.drop(index)

    def display_entries(self):
        return self.entries

    def export_entries(self, filename, export_format):
        if export_format == "xlsx":
            self.entries.to_excel(filename, index=False)
        elif export_format == "pdf":
            self.export_to_pdf(filename)
        elif export_format == "docx":
            self.export_to_docx(filename)
        elif export_format == "txt":
            self.export_to_text(filename)

    def export_to_pdf(self, filename):
        pdf = canvas.Canvas(filename)
        pdf.drawString(100, 800, "Patient Entry List:")
        y_position = 780
        for index, row in self.entries.iterrows():
            y_position -= 15
            formatted_date = self.format_date(row['Entry Time'])
            pdf.drawString(100, y_position, f"{row['Patient Name']}: {formatted_date}")
        pdf.save()

    def export_to_docx(self, filename):
        doc = Document()
        doc.add_heading("Patient Entry List", level=1)
        for index, row in self.entries.iterrows():
            formatted_date = self.format_date(row['Entry Time'])
            doc.add_paragraph(f"{row['Patient Name']}: {formatted_date}")
        doc.save(filename)

    def export_to_text(self, filename):
        with open(filename, 'w') as f:
            for index, row in self.entries.iterrows():
                formatted_date = self.format_date(row['Entry Time'])
                f.write(f"{row['Patient Name']}: {formatted_date}\n")

    def format_date(self, entry_time):
        dt = datetime.datetime.strptime(entry_time, "%Y-%m-%d %H:%M:%S")
        day_of_week = calendar.day_name[dt.weekday()]
        return f"{entry_time} ({day_of_week})"

    def modify_entry_date(self, index, new_date):
        self.entries.at[index, 'Entry Time'] = new_date

    def schedule_follow_up_reminders(self):
        for index, row in self.entries.iterrows():
            patient_name = row['Patient Name']
            for _ in range(30):
                follow_up_time = datetime.datetime.now() + datetime.timedelta(minutes=1)
                print(f"Scheduled follow-up reminder for {patient_name} at {follow_up_time}")
                while datetime.datetime.now() < follow_up_time:
                    time.sleep(1)
                print(f"Reminder for {patient_name}! Current time: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                playsound("path/to/your/sound/file.mp3")  # Replace with the actual path to your sound file

class PatientEntryGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Patient Entry GUI")

        self.patient_list = PatientEntryList()

        self.entry_var = tk.StringVar()
        self.entry_var.set("")

        entry_label = tk.Label(root, text="Enter patient name:")
        entry_label.pack()

        entry_entry = tk.Entry(root, textvariable=self.entry_var)
        entry_entry.pack()

        record_button = tk.Button(root, text="Record Entry", command=self.record_entry)
        record_button.pack()

        edit_button = tk.Button(root, text="Edit Entry", command=self.edit_entry)
        edit_button.pack()

        delete_button = tk.Button(root, text="Delete Entry", command=self.delete_entry)
        delete_button.pack()

        modify_date_button = tk.Button(root, text="Modify Date", command=self.modify_entry_date_dialog)
        modify_date_button.pack()

        export_button = tk.Button(root, text="Export Entries", command=self.export_entries_dialog)
        export_button.pack()

        display_button = tk.Button(root, text="Display Entries", command=self.display_entries)
        display_button.pack()

        self.tree = ttk.Treeview(root, columns=['Patient Name', 'Entry Time'], show='headings')
        self.tree.heading('Patient Name', text='Patient Name')
        self.tree.heading('Entry Time', text='Entry Time')
        self.tree.bind('<ButtonRelease-1>', self.on_tree_click)  # bind click event
        self.tree.pack()

    def record_entry(self):
        patient_name = self.entry_var.get()
        if patient_name:
            self.patient_list.append_entry(patient_name)
            print(f"Entry recorded for {patient_name} at {self.patient_list.entries.loc[self.patient_list.entries['Patient Name'] == patient_name, 'Entry Time'].values[0]}")
            self.display_entries()

    def edit_entry(self):
        selected_item = self.tree.selection()
        if selected_item:
            old_patient_name = self.tree.item(selected_item, 'values')[0]
            new_patient_name = simpledialog.askstring("Edit Entry", f"Edit entry for {old_patient_name} to:")
            if new_patient_name:
                index = int(selected_item[0])
                self.tree.item(selected_item, values=(new_patient_name, datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
                self.patient_list.entries.at[index, 'Patient Name'] = new_patient_name

    def delete_entry(self):
        selected_item = self.tree.selection()
        if selected_item:
            index = int(selected_item[0])
            self.tree.delete(selected_item)
            self.patient_list.delete_entry(index)

    def modify_entry_date_dialog(self):
        selected_item = self.tree.selection()
        if selected_item:
            index = int(selected_item[0])
            old_date = self.patient_list.entries.at[index, 'Entry Time']
            new_date = self.modify_entry_date_input_dialog(old_date)
            if new_date:
                self.patient_list.modify_entry_date(index, new_date)
                self.display_entries()

    def modify_entry_date_input_dialog(self, old_date):
        new_date = simpledialog.askstring("Modify Date", "Enter new date and time (YYYY-MM-DD HH:MM:SS):", initialvalue=old_date)
        return new_date

    def export_entries_dialog(self):
        export_format = simpledialog.askstring("Export Format", "Choose export format (xlsx, pdf, docx, txt):").lower()
        if export_format in ["xlsx", "pdf", "docx", "txt"]:
            filename = filedialog.asksaveasfilename(defaultextension=f".{export_format}", filetypes=[(f"{export_format.upper()} files", f"*.{export_format}")])
            if filename:
                self.patient_list.export_entries(filename, export_format)
                print(f"Entries exported to {filename}")
        else:
            print("Invalid export format. Choose xlsx, pdf, docx, or txt.")

    def display_entries(self):
        entries = self.patient_list.display_entries()
        for i in self.tree.get_children():
            self.tree.delete(i)
        for index, row in entries.iterrows():
            self.tree.insert('', index, index, values=(row['Patient Name'], self.patient_list.format_date(row['Entry Time'])))

    def on_tree_click(self, event):
        pass  # You can remove this function if not needed

if __name__ == "__main__":
    root = tk.Tk()
    app = PatientEntryGUI(root)
    root.mainloop()
